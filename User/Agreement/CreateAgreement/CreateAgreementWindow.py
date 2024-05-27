import sys
from PyQt5.QtWidgets import QApplication, QMainWindow,QCheckBox, QVBoxLayout, QMessageBox, QFileDialog
from WindowsPY.CreateAgreement import Ui_CreateAgreement
import requests
import os
from pathlib import Path
import User.Agreement.CreateAgreement.AllTPAGreementWindow as m
from WindowsPY.Admin.CreateGost import Ui_CreateGOST
from PyQt5.QtGui import QIcon
from User.CreateTp.functions import getGOST
from docx import Document
import re
from WindowSet import WINDOW_HEIGHT, WINDOW_WIDTH, center_window
class CreateAgreement(QMainWindow, Ui_CreateAgreement):
    def __init__(self, parent=None, UserData = {}, icon = QIcon(''), dataThisTP = {}):
        super().__init__(parent)
        self.setupUi(self)
        self.userD = UserData
        self.icon = icon
        self.setFixedSize(WINDOW_WIDTH, WINDOW_HEIGHT)
        center_window(self)
        self.dataThisTP = dataThisTP
        self.pushButtonBack.clicked.connect(self.go_back)
        self.pushButtonCheckOldAgreement_4.clicked.connect(self.selectFile)
        self.pushButtonDownload.clicked.connect(self.downloadTP)
        self.pushButtonToAgreemennt.clicked.connect(self.agreement)
        self.initUI()

    def initUI(self):
        self.checkboxes = []
        self.GOSTS = getGOST()
        self.GOSTTP = requests.get('http://127.0.0.1:8000/Связь%20ГОСТ%20и%20ТП/').json()
        self.verticalLayout = QVBoxLayout(self.scrollAreaWidgetContents_2)
        for i in range(len(self.GOSTS)):
            isTRUE = 0
            #checkbox = QCheckBox(f"Элемент {i+1}", self.scrollAreaWidgetContents_2)
            checkbox = QCheckBox(self.GOSTS[i]['gostName'], self.scrollAreaWidgetContents_2)
            checkbox.setObjectName(f"checkBox_{i}")

            for j in range(len(self.GOSTTP)):
                if self.GOSTTP[j]['idDOCK'] == self.dataThisTP['id']:
                    if self.GOSTTP[j]['idGOST'] == self.GOSTS[i]['id']:
                        isTRUE = 1
                        break
            
            if (isTRUE == 1):
                checkbox.setChecked(True)
            else:
                checkbox.setChecked(False)
            isTRUE = 0
            self.checkboxes.append(checkbox)
            self.verticalLayout.addWidget(checkbox)
        
        self.labelName.setText(self.dataThisTP['TpName'])


    def selectFile(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Выбрать файл")
        if file_path:
            self.pathToFile = Path(file_path)     
            self.label_3.setText(file_path)
        else:
            return
        
        doc = Document(self.pathToFile)
        found_gosts = []
        pattern = r'ГОСТ [\w-]+'
        matches = []

        for paragraph in doc.paragraphs:
            if re.search(pattern, paragraph.text):
                matches.extend(re.findall(pattern, paragraph.text))
            for run in paragraph.runs:
                for gost in self.GOSTS:
                    if gost['gostName'] in run.text:
                        found_gosts.append(gost)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = ' '.join([paragraph.text for paragraph in cell.paragraphs]) + ' ' + cell.text
                    if re.search(pattern, cell_text):
                        matches.extend(re.findall(pattern, cell_text))
                        for gost in self.GOSTS:
                            if gost['gostName'] in cell_text:
                                found_gosts.append(gost)

        for checkbox in self.checkboxes:
            checkbox.setChecked(False)
        if len(found_gosts) == 0:
            QMessageBox.information(self.centralwidget, "Ошибка", "В документе не указаны ГОСТ")
            return

        for i in range (len(found_gosts)):
            for checkbox in range(len(self.checkboxes)):
                if self.checkboxes[checkbox].text()== found_gosts[i]['gostName']:
                    self.checkboxes[checkbox].setChecked(True)
            
        matches = [match for match in matches if match not in [gost['gostName'] for gost in found_gosts]]
        if len(matches) > 0:
            error_message = "В документе указаны ГОСТ которых нет в системе:\n" + "\n".join(matches) + "\nОбратитесь к администратору для их добавления"
            QMessageBox.information(self.centralwidget, "Ошибка", error_message)




    def downloadTP(self):
        url = self.dataThisTP['currentVersionTP']
        save_path = self.lineEditPathToCurrentVersion.text()
        if len(save_path) < 2:
            QMessageBox.information(self.centralwidget, "Ошибка", "Укажите корректный путь")
            return
        if url is None:
            return
        response = requests.get(url)
        if response.status_code == 200:
            file_path = os.path.join(save_path, 'ТП ' + self.dataThisTP['TpName']+ ".docx")
            if len(file_path) > 3:
                try:
                    with open(file_path, "wb") as file:
                        file.write(response.content)
                        file.close()
                        QMessageBox.information(self.centralwidget, "Успешно", "Файл успешно загружен")
                        return
                except:
                    pass
        QMessageBox.information(self.centralwidget, "Ошибка", "Ошибка Загрузки!!")
    
    def agreement(self):
        self.ALLAGrementt = requests.get('http://127.0.0.1:8000/Согласование%20ТП/').json()
        for i in range(len(self.ALLAGrementt)):
            if (self.ALLAGrementt[i]['idTP'] == self.dataThisTP['id'] and self.ALLAGrementt[i]['isActual'] == True):
                QMessageBox.information(self.centralwidget, "Ошибка", "Согласование данного ГОСТ уже существует")
                return
        url = "http://127.0.0.1:8000/Согласование%20ТП/"
        self.selected_items = [checkbox.text() for checkbox in self.checkboxes if checkbox.isChecked()]
        #self.selected_items
        self.GOSTSID = ''
        for i in range(len(self.selected_items)):
            for j in range(len(self.GOSTS)):
                if self.GOSTS[j]['gostName'] == self.selected_items[i]:
                    self.GOSTSID += str(self.GOSTS[j]['id']) + ' '
        self.GOSTSID += '-'
        if (len(self.GOSTSID) < 2):
            QMessageBox.information(self.centralwidget, "Ошибка", "Укажите ГОСТ")
            return
        try:
            with open(self.pathToFile, 'rb') as file:
                file_data = file.read()
        except:
            QMessageBox.information(self.centralwidget, "Ошибка", "Укажите файл")
            return

        files = {
            'dock': ('TP', file_data, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }

        data = {
            "idTP": self.dataThisTP['id'],
            "creator": self.userD['id'],
            "isActual": True,
            "idTpStringNew": self.GOSTSID,
            "NewName": '-' if len(self.lineEditNewName.text()) < 1 else self.lineEditNewName.text(),
            "IsNewTP": False,
            "comment": '-' if len(self.lineEditComment.text()) < 1 else len(self.lineEditComment.text()) 
        }
        r = requests.post(url, files=files, data=data)
        try:
            if r.status_code == 201:
                QMessageBox.information(self.centralwidget, "Успешно", "Отправлено на согласование")
            else:
                QMessageBox.information(self.centralwidget, "Ошибка", r.status_code + r.text)
        except:
            QMessageBox.information(self.centralwidget, "Ошибка", "Ошибка в запросе")

        files = {
            'dock': ('newDockVersion', file_data, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        url = f"http://127.0.0.1:8000/ТП/{self.dataThisTP['id']}/"
        r = requests.post(url, files=files)
    def go_back(self):
        self.menu = m.CreateAgreementALLTP(UserData=self.userD, icon=self.icon)
        self.menu.setWindowIcon(self.icon) 
        self.menu.show()
        self.close()

