import sys
from PyQt5.QtWidgets import QApplication, QMainWindow,QCheckBox, QVBoxLayout, QMessageBox, QFileDialog
from WindowsPY.CreateTP import Ui_CreateTp
from User.CreateTp.functions import getGOST
import requests
import os
from pathlib import Path
import Adminestrator.GOSTS.GOSTSWindow as m
from WindowsPY.Admin.ChangeGost import Ui_ChangeGost
from PyQt5.QtGui import QIcon
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from WindowSet import WINDOW_HEIGHT, WINDOW_WIDTH, center_window
import re
from docx.enum.text import WD_COLOR_INDEX
class ChangeGost(QMainWindow, Ui_ChangeGost):
    def __init__(self, parent=None, UserData = {}, icon = QIcon(''), dataGost = {}):
        super().__init__(parent)
        self.setupUi(self)
        self.userD = UserData
        self.icon = icon
        self.dataGost = dataGost
        self.initUI()
        self.setFixedSize(WINDOW_WIDTH, WINDOW_HEIGHT)
        center_window(self)


    def initUI(self):
        self.pushButtonBack.clicked.connect(self.go_back)
        self.pushButtonSelectFile.clicked.connect(self.selectFile)
        self.pushButtonDownloadTp.clicked.connect(self.downloadGOST)
        self.pushButtonAgreement.clicked.connect(self.changeGOST)
        self.pushButtonDeleteGost.clicked.connect(self.deleteGost)
        
        self.lineEditName.setText(self.dataGost["gostName"])

    def selectFile(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Выбрать файл")
        if file_path:
            self.pathToFile = Path(file_path)     
        self.labelPathToFile.setText(file_path)

    def downloadGOST(self):
        url = self.dataGost['file']
        save_path = self.lineEditPathDownload.text()
        if url is None:
            return
        if len(save_path) < 2:
            QMessageBox.information(self.centralwidget, "Ошибка", "Укажите верный путь")
            return

        response = requests.get(url)

        if response.status_code == 200:
            file_path = os.path.join(save_path, str(self.dataGost["gostName"]) + ".pdf")
            if len(file_path) > 2:
                try:
                    with open(file_path, "wb") as file:
                        file.write(response.content)
                        file.close()
                        QMessageBox.information(self.centralwidget, "Успешно", "Файл успешно загружен")
                        return
                except:
                    pass
        QMessageBox.information(self.centralwidget, "Ошибка", "Ошибка Загрузки!!")




    def deleteGost(self):
        try:
            r = requests.delete(f'http://127.0.0.1:8000/ГОСТ/{self.dataGost['id']}') 
            QMessageBox.information(self.centralwidget, "Успешно", "ГОСТ удалён")
        except:
            QMessageBox.information(self.centralwidget, "Ошибка", "Данного ГОСТ не существует")

    def changeGOST(self):
        url = f"http://127.0.0.1:8000/ГОСТ/{self.dataGost['id']}/"
        flagFile = 0
        try:
            with open(self.pathToFile, 'rb') as file:
                file_data = file.read()
                flagFile = 1
        except:
            flagFile = 0

        if (self.lineEditName):
            pass   
        else:
            QMessageBox.information(self.centralwidget, "Ошибка", "Укажите наименование ГОСТ")
            return
        if (flagFile == 1):
            files = {
                'file': (self.lineEditName.text(), file_data, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            }
        data = {
            "gostName": self.lineEditName.text(),
            "idCreator": self.userD['id']
        }
        if(flagFile == 1):
            r = requests.patch(url, files=files, data=data)
        else:
            r = requests.patch(url, data=data)
        try:
            if r.status_code == 200:
                QMessageBox.information(self.centralwidget, "Успешно", "ГОСТ изменён")
            else:
                QMessageBox.information(self.centralwidget, "Ошибка", r.status_code + r.text)
        except:
            print(r.status_code)
            print(r.text)
            QMessageBox.information(self.centralwidget, "Ошибка", "Ошибка в запросе")

        #изменение необходимости изменять ТП
        gots_TP = requests.get(f'http://127.0.0.1:8000/Связь%20ГОСТ%20и%20ТП/') 
        TP = requests.get(f'http://127.0.0.1:8000/ТП/') 

        for i in range(len(gots_TP.json())):
            if gots_TP.json()[i]['idGOST'] == self.dataGost['id']:
                url = f"http://127.0.0.1:8000/ТП/{gots_TP.json()[i]['idDOCK']}/"
                data = {
                "needForChange": True
                }
                r = requests.patch(url, data=data)

            if gots_TP.json()[i]['idGOST'] == self.dataGost['id']:    
                urlTPChange = f'http://127.0.0.1:8000/ТП/{gots_TP.json()[i]["idDOCK"]}/'
                dataThisGost = requests.get(urlTPChange).json()['currentVersionTP']
                dockiment = requests.get(dataThisGost)
                file_path = os.path.join(os.getcwd(), "TEMP.docx")
                with open(file_path, "wb") as file:
                        file.write(dockiment.content)
                        file.close()
                doc = Document(file_path)
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        if self.dataGost['gostName'] in run.text:
                            run.font.highlight_color = WD_COLOR_INDEX.RED 
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell_text = ' '.join([paragraph.text for paragraph in cell.paragraphs]) + ' ' + cell.text
                            if re.search(self.dataGost['gostName'], cell_text):
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        print(run.text)
                                        if self.dataGost['gostName'] in run.text:
                                            run.font.highlight_color = WD_COLOR_INDEX.RED
                    


                doc.save(file_path)
                url = f"http://127.0.0.1:8000/ТП/{gots_TP.json()[i]['idDOCK']}/"
                files = {'currentVersionTP': open(file_path, 'rb')}
                response = requests.patch(url, files=files)



        #files['currentVersionTP'].close()
        #os.remove(file_path)



    def go_back(self):
        #self.close()
        #self.parent().show()  
        self.allgostss = m.ALLGosts(UserData=self.userD, icon=self.icon)
        self.allgostss.setWindowIcon(self.icon) 
        self.allgostss.show()
        self.close()
